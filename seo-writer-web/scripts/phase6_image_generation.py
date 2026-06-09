#!/usr/bin/env python3
"""Generate planned article images, then optionally insert approved outputs into Word."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches

from image_providers import get_provider
from phase5_5_image_planning import ensure_plan_image_count


REQUIRED_IMAGE_FIELDS = ("id", "type", "prompt", "alt_text", "insertion_marker")
VOLCENGINE_ARK_IMAGE_PROVIDERS = {"doubao", "volcengine_ark"}
VOLCENGINE_ARK_IMAGE_MODEL_IDS_BY_DISPLAY = {
    "Doubao-Seedream-5.0-lite": "doubao-seedream-5-0-lite-260128",
    "Doubao-Seedream-5.0": "doubao-seedream-5-0-260128",
    "Doubao-Seedream-4.5": "doubao-seedream-4-5-251128",
    "Doubao-Seedream-4.0": "doubao-seedream-4-0-250828",
}
VOLCENGINE_MODEL_ID_HINT = (
    "当前 final_model_for_request 看起来像 UI 展示名称。请使用火山方舟官方 Model ID / Endpoint ID，"
    "例如 doubao-seedream-5-0-lite-260128，而不是 Doubao-Seedream-5.0-lite。"
)
INVALID_ENDPOINT_OR_MODEL_MESSAGE = (
    "当前传入的 image_model_id 或 image_endpoint_id 不存在，或账号没有权限。"
    "请确认模型服务已开通，并从火山方舟官方文档或控制台复制真实 Model ID / Endpoint ID，而不是 UI 展示名称。"
)
INVALID_VOLCENGINE_IMAGE_SIZE_MESSAGE = (
    "当前传入的图片 size 不符合火山方舟生图要求。Seedream 图片尺寸需至少 3686400 像素；"
    "系统已按比例使用 2K 级尺寸，例如 16:9 使用 2560x1440。请重试 Phase 6。"
)
INCOMPLETE_IMAGE_RESPONSE_MESSAGE = (
    "图片接口响应在传输过程中中断，导致图片数据没有读完整。系统已在底层自动重试；"
    "如果仍失败，请在 Phase 6 点击重试，已成功的图片会保留，只会补生成失败的部分。"
)


def now() -> str:
    return datetime.now(timezone.utc).isoformat()


class Phase6Log:
    def __init__(self, path: Path) -> None:
        self.path = path
        self.lines: list[str] = []

    def write(self, message: str) -> None:
        self.lines.append(f"[{now()}] {message}")

    def save(self) -> None:
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.path.write_text("\n".join(self.lines) + "\n", encoding="utf-8")


class ImageModelConfigurationError(ValueError):
    pass


def read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8-sig"))


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def approved_image_ids_from_review(path: Path) -> set[str]:
    review = read_json(path, {})
    images = review.get("images", {}) if isinstance(review, dict) else {}
    if not isinstance(images, dict):
        return set()
    return {
        str(image_id).strip()
        for image_id, item in images.items()
        if str(image_id).strip()
        and isinstance(item, dict)
        and item.get("status") == "approved"
    }


def copy_original_docx(final_docx: Path, output_docx: Path, log: Phase6Log, reason: str) -> None:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    if final_docx.exists():
        shutil.copyfile(final_docx, output_docx)
        log.write(f"Copied original Word output to {output_docx}. Reason: {reason}")
    else:
        log.write(f"Original Word file is missing; cannot copy fallback output: {final_docx}")


def policy_ready(plan: dict[str, Any]) -> tuple[bool, str]:
    policy = plan.get("policy_check")
    if not isinstance(policy, dict):
        return False, "image_plan.json missing policy_check"
    if policy.get("ready_for_generation") is not True:
        return False, "policy_check.ready_for_generation is not true"
    failed = [key for key, value in policy.items() if isinstance(value, bool) and value is False]
    if failed:
        return False, f"policy_check failed: {', '.join(failed)}"
    return True, ""


def normalize_image(image: dict[str, Any], index: int) -> dict[str, Any]:
    image_id = str(image.get("id") or f"IMAGE_{index}").strip()
    match = re.search(r"(\d+)", image_id)
    if match and not image_id.startswith("IMAGE_"):
        image_id = f"IMAGE_{match.group(1)}"
    normalized = dict(image)
    normalized["id"] = image_id
    normalized["type"] = str(image.get("type") or "").strip()
    normalized["prompt"] = str(image.get("prompt") or "").strip()
    normalized["alt_text"] = str(image.get("alt_text") or "").strip()
    normalized["insertion_marker"] = f"[{image_id}]"
    normalized["insert_after_text"] = str(image.get("insert_after_text") or "").strip()
    normalized["insert_before_heading"] = str(image.get("insert_before_heading") or "").strip()
    return normalized


def validate_plan_images(plan: dict[str, Any], log: Phase6Log) -> list[dict[str, Any]]:
    images = plan.get("images")
    if not isinstance(images, list):
        log.write("image_plan.json images is missing or not a list")
        return []

    valid: list[dict[str, Any]] = []
    for index, raw in enumerate(images, start=1):
        if not isinstance(raw, dict):
            continue
        image = normalize_image(raw, index)
        missing = [field for field in REQUIRED_IMAGE_FIELDS if not image.get(field)]
        if missing:
            log.write(f"{image.get('id', 'UNKNOWN')} missing required fields: {', '.join(missing)}")
            continue
        valid.append(image)
    return valid


def config_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    return str(value or "").strip().lower() in {"1", "true", "yes", "on"}


def resolve_model_request_fields(config: dict[str, Any]) -> dict[str, Any]:
    provider = str(config.get("provider") or "volcengine_ark").strip()
    provider_normalized = provider.lower()
    model_display_name = str(
        config.get("image_model_display_name")
        or config.get("model_display_name")
        or config.get("model_name")
        or ""
    ).strip()
    image_model_id = str(config.get("image_model_id") or config.get("model_id") or "").strip()
    image_endpoint_id = str(config.get("image_endpoint_id") or config.get("endpoint_id") or "").strip()
    use_endpoint_id = config_bool(config.get("use_endpoint_id"))
    api_type = str(config.get("api_type") or "").strip()
    if not api_type:
        api_type = "volcengine_ark_images_generations" if provider_normalized in VOLCENGINE_ARK_IMAGE_PROVIDERS else f"{provider_normalized}_images"

    if provider_normalized in VOLCENGINE_ARK_IMAGE_PROVIDERS:
        if image_model_id in VOLCENGINE_ARK_IMAGE_MODEL_IDS_BY_DISPLAY:
            image_model_id = VOLCENGINE_ARK_IMAGE_MODEL_IDS_BY_DISPLAY[image_model_id]
        elif not image_model_id and model_display_name in VOLCENGINE_ARK_IMAGE_MODEL_IDS_BY_DISPLAY:
            image_model_id = VOLCENGINE_ARK_IMAGE_MODEL_IDS_BY_DISPLAY[model_display_name]
    elif not image_model_id:
        image_model_id = str(config.get("model") or "").strip()

    final_model = image_endpoint_id if use_endpoint_id else image_model_id

    return {
        "image_provider": provider,
        "provider_normalized": provider_normalized,
        "api_type": api_type,
        "image_model_display_name": model_display_name,
        "image_model_id": image_model_id,
        "image_endpoint_id": image_endpoint_id,
        "use_endpoint_id": use_endpoint_id,
        "final_model_for_request": final_model,
    }


def validate_model_request_fields(fields: dict[str, Any]) -> None:
    final_model = str(fields.get("final_model_for_request") or "").strip()
    display_name = str(fields.get("image_model_display_name") or "").strip()
    provider = str(fields.get("provider_normalized") or "").strip()
    if not final_model:
        raise ImageModelConfigurationError(
            "final_model_for_request is empty. Set image_model_id when use_endpoint_id=false, "
            "or image_endpoint_id when use_endpoint_id=true."
        )
    if provider in VOLCENGINE_ARK_IMAGE_PROVIDERS:
        if display_name and final_model == display_name:
            raise ImageModelConfigurationError(VOLCENGINE_MODEL_ID_HINT)
        if re.search(r"[A-Z.\s]", final_model):
            raise ImageModelConfigurationError(VOLCENGINE_MODEL_ID_HINT)


def normalize_config_model_fields(config: dict[str, Any], fields: dict[str, Any]) -> None:
    image_model_display_name = str(fields.get("image_model_display_name") or "")
    image_model_id = str(fields.get("image_model_id") or "")
    image_endpoint_id = str(fields.get("image_endpoint_id") or "")
    final_model = str(fields.get("final_model_for_request") or "")
    config["provider"] = str(fields.get("image_provider") or config.get("provider") or "")
    config["api_type"] = str(fields.get("api_type") or config.get("api_type") or "")
    config["model_display_name"] = image_model_display_name
    config["image_model_display_name"] = image_model_display_name
    config["model_id"] = image_model_id
    config["image_model_id"] = image_model_id
    config["endpoint_id"] = image_endpoint_id
    config["image_endpoint_id"] = image_endpoint_id
    config["use_endpoint_id"] = bool(fields.get("use_endpoint_id", False))
    config["model"] = final_model
    config["final_model_for_request"] = final_model


def log_model_request_fields(log: Phase6Log, fields: dict[str, Any]) -> None:
    log.write(f"image_provider={fields.get('image_provider', '')}")
    log.write(f"api_type={fields.get('api_type', '')}")
    log.write(f"image_model_display_name={fields.get('image_model_display_name', '')}")
    log.write(f"image_model_id={fields.get('image_model_id', '')}")
    log.write(f"image_endpoint_id={fields.get('image_endpoint_id', '')}")
    log.write(f"use_endpoint_id={fields.get('use_endpoint_id', False)}")
    log.write(f"final_model_for_request={fields.get('final_model_for_request', '')}")


def normalize_image_generation_error(error: Exception) -> str:
    message = str(error)
    if "InvalidEndpointOrModel.NotFound" in message:
        return INVALID_ENDPOINT_OR_MODEL_MESSAGE
    if "image size must be at least 3686400 pixels" in message:
        return INVALID_VOLCENGINE_IMAGE_SIZE_MESSAGE
    if "IncompleteRead" in message or "response was interrupted" in message:
        return INCOMPLETE_IMAGE_RESPONSE_MESSAGE
    return message


def image_file_openable(path: Path) -> bool:
    if not path.exists() or path.stat().st_size == 0:
        return False
    header = path.read_bytes()[:16]
    if header.startswith(b"\x89PNG\r\n\x1a\n"):
        return True
    if header.startswith(b"\xff\xd8\xff"):
        return True
    if header.startswith(b"RIFF") and b"WEBP" in header:
        return True
    try:
        from PIL import Image  # type: ignore

        with Image.open(path) as image:
            image.verify()
        return True
    except Exception:
        return False


def prompt_contains_any(prompt: str, needles: list[str]) -> bool:
    lowered = prompt.lower()
    return any(needle in lowered for needle in needles)


def compliance_for_image(image: dict[str, Any], file_path: Path, plan: dict[str, Any]) -> dict[str, Any]:
    prompt = str(image.get("prompt") or "")
    lowered = prompt.lower()
    image_type = str(image.get("type") or "").lower()
    image_id = str(image.get("id") or "")
    notes = image.get("policy_notes") if isinstance(image.get("policy_notes"), dict) else {}
    photorealistic_types = {"photo-hero-composite", "photo-ui-composite", "photo-learning-scene"}
    allowed_types = photorealistic_types | {"infographic", "data chart", "ui-screenshot-modified", "scene-illustration"}
    requires_photo_style = image_type in photorealistic_types
    includes_generated_person = (
        not requires_photo_style
        or (
            ("generated" in lowered or "non-identifiable" in lowered or "non identifiable" in lowered)
            and ("learner" in lowered or "person" in lowered or "adult" in lowered)
        )
    )
    checks = {
        "file_exists": file_path.exists() and file_path.stat().st_size > 0,
        "file_openable": image_file_openable(file_path),
        "type_matches_plan": image_type in allowed_types,
        "hero_is_photorealistic_brand_composite": image_id != "IMAGE_1" or image_type in photorealistic_types,
        "uses_generated_non_identifiable_person": includes_generated_person,
        "no_large_visible_text": "large visible text" not in lowered or "no large visible text" in lowered,
        "not_stock_photo_style": not prompt_contains_any(
            prompt,
            ["generic lifestyle photo", "stock-photo-style", "stock photo style", "young adult wearing headphones"],
        ),
        "uses_brand_style": "#5856d6" in lowered or "hellotalk purple" in lowered,
        "has_filled_cards": ("filled" in lowered or "rounded" in lowered) and ("card" in lowered or "block" in lowered or "chip" in lowered),
        "has_pill_badges_or_containers": "pill" in lowered or "visual container" in lowered or "container" in lowered,
        "has_decorative_geometry": "geometry" in lowered or "geometric" in lowered,
        "no_fake_features": bool(notes.get("no_fake_features", True)),
        "not_near_hellotalk_mention": bool(notes.get("not_near_hellotalk_mention", True)),
        "not_before_final_cta": bool(notes.get("not_before_final_cta", True)),
    }

    if image_type == "ui-screenshot-modified":
        checks["ui_screenshot_source_valid"] = bool(image.get("source_ui"))
    else:
        checks["ui_screenshot_source_valid"] = True

    critical_keys = [
        "file_exists",
        "file_openable",
        "type_matches_plan",
        "hero_is_photorealistic_brand_composite",
        "uses_generated_non_identifiable_person",
        "not_stock_photo_style",
        "uses_brand_style",
        "has_filled_cards",
        "has_pill_badges_or_containers",
        "has_decorative_geometry",
        "no_fake_features",
        "not_near_hellotalk_mention",
        "not_before_final_cta",
        "ui_screenshot_source_valid",
    ]
    failed = [key for key in critical_keys if not checks.get(key)]
    return {
        "id": image_id,
        "type": image.get("type", ""),
        "checked_by": "deterministic file, prompt, and image_plan policy checks",
        "checks": checks,
        "compliant": not failed,
        "failed_checks": failed,
    }


def existing_success_by_id(existing_metadata: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    successes: dict[str, dict[str, Any]] = {}
    for entry in existing_metadata:
        image_id = str(entry.get("id") or "").strip()
        file_path = Path(str(entry.get("file_path") or ""))
        if image_id and entry.get("status") == "success" and file_path.exists() and image_file_openable(file_path):
            successes[image_id] = entry
    return successes


def successful_generated_count(metadata: list[dict[str, Any]]) -> int:
    count = 0
    for entry in metadata:
        file_path = Path(str(entry.get("file_path") or ""))
        if entry.get("status") == "success" and file_path.exists() and image_file_openable(file_path):
            count += 1
    return count


def insertable_generated_count(metadata: list[dict[str, Any]], approved_image_ids: set[str]) -> int:
    count = 0
    for entry in metadata:
        image_id = str(entry.get("id") or "").strip()
        file_path = Path(str(entry.get("file_path") or ""))
        if not file_path.exists() or not image_file_openable(file_path):
            continue
        if entry.get("status") == "success" or image_id in approved_image_ids:
            count += 1
    return count


def remove_stale_output_docx(output_docx: Path, output_dir: Path, log: Phase6Log, reason: str) -> None:
    try:
        resolved_output = output_docx.resolve()
        resolved_dir = output_dir.resolve()
        if output_docx.exists() and resolved_dir in resolved_output.parents:
            output_docx.unlink()
            log.write(f"Removed stale final Word output. Reason: {reason}")
    except Exception as exc:
        log.write(f"Could not remove stale final Word output: {exc}")


def generate_images(
    images: list[dict[str, Any]],
    plan: dict[str, Any],
    config: dict[str, Any],
    images_dir: Path,
    log: Phase6Log,
    existing_metadata: list[dict[str, Any]] | None = None,
    force_regenerate_ids: set[str] | None = None,
    review_comment: str = "",
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    model_fields = resolve_model_request_fields(config)
    provider_name = str(model_fields["image_provider"] or "volcengine_ark")
    model = str(model_fields["final_model_for_request"])
    display_model = str(model_fields["image_model_display_name"])
    base_url = str(config.get("base_url") or "")
    output_format = str(config.get("output_format") or "png").lstrip(".")
    aspect_ratio = str(config.get("default_aspect_ratio") or "16:9")
    retry_count = int(config.get("retry_count") or 0)
    allow_non_compliant = bool(config.get("allow_non_compliant_images", False))
    provider = get_provider(provider_name, base_url)
    metadata: list[dict[str, Any]] = []
    compliance_report: list[dict[str, Any]] = []
    existing_successes = existing_success_by_id(existing_metadata or [])
    force_ids = force_regenerate_ids or set()

    for image in images:
        image_id = str(image["id"])
        output_path = images_dir / f"{image_id}.{output_format}"
        base_prompt = str(image.get("prompt", ""))
        prompt_for_request = base_prompt
        if image_id in force_ids and review_comment.strip():
            prompt_for_request = (
                f"{base_prompt}\n\n"
                "Human image review revision request:\n"
                f"{review_comment.strip()}\n\n"
                "Regenerate only this planned image. Keep the original article context, placement, "
                "alt text intent, style policy, and compliance constraints unchanged."
            )
        entry = {
            "id": image_id,
            "provider": provider_name,
            "api_type": str(model_fields["api_type"]),
            "image_model_display_name": display_model,
            "image_model_id": str(model_fields["image_model_id"]),
            "image_endpoint_id": str(model_fields["image_endpoint_id"]),
            "use_endpoint_id": bool(model_fields["use_endpoint_id"]),
            "final_model_for_request": model,
            "model": display_model,
            "model_id": str(model_fields["image_model_id"]),
            "endpoint_id": str(model_fields["image_endpoint_id"]),
            "effective_model": model,
            "prompt": prompt_for_request,
            "original_prompt": base_prompt,
            "review_comment": review_comment.strip() if image_id in force_ids else "",
            "alt_text": image.get("alt_text", ""),
            "type": image.get("type", ""),
            "placeholder": image.get("insertion_marker", f"[{image_id}]"),
            "insert_after_text": image.get("insert_after_text", ""),
            "insert_before_heading": image.get("insert_before_heading", ""),
            "file_path": str(output_path),
            "status": "failed",
            "retry_count": 0,
            "compliant": False,
        }
        log.write(f"Starting image generation for {image_id}")
        existing_success = existing_successes.get(image_id)
        if existing_success and image_id not in force_ids:
            existing_path = Path(str(existing_success.get("file_path") or ""))
            compliance = compliance_for_image(image, existing_path, plan)
            compliance_report.append(compliance)
            if bool(compliance.get("compliant")) or allow_non_compliant:
                entry["file_path"] = str(existing_path)
                entry["status"] = "success"
                entry["retry_count"] = int(existing_success.get("retry_count") or 0)
                entry["compliant"] = bool(compliance.get("compliant"))
                entry["compliance_failed_checks"] = compliance.get("failed_checks", [])
                entry["reused_existing_success"] = True
                metadata.append(entry)
                log.write(f"{image_id} reused existing successful image: {existing_path}")
                continue
            log.write(f"{image_id} existing image is no longer compliant; regenerating.")

        if not model:
            entry["error"] = "Missing image model, model_id, or endpoint_id in project configuration."
            metadata.append(entry)
            log.write(f"{image_id} failed: {entry['error']}")
            continue

        for attempt in range(retry_count + 1):
            entry["retry_count"] = attempt
            try:
                generated = provider.generate_image(
                    prompt=prompt_for_request,
                    model=model,
                    output_path=output_path,
                    aspect_ratio=aspect_ratio,
                    output_format=output_format,
                )
                entry["file_path"] = str(generated)
                compliance = compliance_for_image(image, Path(generated), plan)
                compliance_report.append(compliance)
                entry["compliant"] = bool(compliance.get("compliant"))
                entry["compliance_failed_checks"] = compliance.get("failed_checks", [])
                if entry["compliant"] or allow_non_compliant:
                    entry["status"] = "success"
                    entry.pop("error", None)
                    log.write(f"{image_id} generated successfully: {generated}")
                    log.write(f"{image_id} compliance={entry['compliant']} failed_checks={entry['compliance_failed_checks']}")
                    break
                entry["status"] = "non_compliant"
                entry["error"] = "Generated image failed compliance checks."
                log.write(f"{image_id} compliance failed on attempt {attempt + 1}: {entry['compliance_failed_checks']}")
            except Exception as exc:
                error_message = normalize_image_generation_error(exc)
                entry["error"] = error_message
                log.write(f"{image_id} attempt {attempt + 1}/{retry_count + 1} failed: {error_message}")
        metadata.append(entry)
    return metadata, compliance_report


def seo_skill_dir() -> Path:
    configured = os.environ.get("SEO_SKILL_DIR", "../skills/seo-article-writer")
    path = Path(configured)
    if path.is_absolute():
        return path
    return (Path.cwd() / path).resolve()


def convert_slots_markdown_to_docx(article_md: Path, fallback_docx: Path, output_dir: Path, log: Phase6Log) -> Path:
    if not article_md.exists():
        log.write(f"slot article markdown is missing: {article_md}")
        return fallback_docx
    script = seo_skill_dir() / "scripts" / "generate_docx.py"
    if not script.exists():
        log.write(f"generate_docx.py not found for slot markdown conversion: {script}")
        return fallback_docx
    slot_docx = output_dir / ".cache" / "final_article_with_image_slots.docx"
    slot_docx.parent.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [sys.executable, str(script), str(article_md), str(slot_docx)],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode != 0 or not slot_docx.exists():
        log.write(f"slot markdown conversion failed: {(result.stderr or result.stdout).strip()}")
        return fallback_docx
    log.write(f"slot markdown converted to Word: {slot_docx}")
    return slot_docx


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_element)
    new_paragraph = Paragraph(new_paragraph_element, paragraph._parent)
    new_paragraph.add_run(text)
    return new_paragraph


def image_caption_text(entry: dict[str, Any]) -> str:
    description = str(
        entry.get("alt_text")
        or entry.get("description")
        or entry.get("visual_concept")
        or ""
    ).strip()
    prompt = str(entry.get("prompt") or "").strip()
    return f"[{description}]（{prompt}）"


def insert_image_in_paragraph(
    paragraph,
    image_path: Path,
    caption_text: str,
    width_inches: float = 5.8,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    if caption_text:
        insert_paragraph_after(paragraph, caption_text)


def append_generated_image(document: Document, image_path: Path, caption_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    if caption_text:
        document.add_paragraph(caption_text)


def remove_leftover_placeholders(document: Document) -> None:
    for paragraph in iter_paragraphs(document):
        if re.search(r"\[IMAGE_\d+\]", paragraph.text):
            clear_paragraph(paragraph)


def insert_images_into_docx(
    source_docx: Path,
    original_docx: Path,
    output_docx: Path,
    metadata: list[dict[str, Any]],
    allow_non_compliant: bool,
    approved_image_ids: set[str],
    log: Phase6Log,
) -> None:
    successful = [
        entry
        for entry in metadata
        if (
            (
                entry.get("status") == "success"
                and (entry.get("compliant") or allow_non_compliant)
            )
            or str(entry.get("id") or "").strip() in approved_image_ids
        )
        and Path(str(entry.get("file_path", ""))).exists()
    ]
    if not successful:
        copy_original_docx(original_docx, output_docx, log, "no compliant images were generated successfully")
        return

    document = Document(source_docx)
    appended_heading_added = False

    for entry in successful:
        image_path = Path(str(entry["file_path"]))
        if str(entry.get("id") or "").strip() in approved_image_ids:
            entry["approved_by_human_review"] = True
        placeholder = str(entry.get("placeholder") or f"[{entry['id']}]")
        caption_text = image_caption_text(entry)
        inserted = False

        for paragraph in iter_paragraphs(document):
            if placeholder in paragraph.text:
                insert_image_in_paragraph(paragraph, image_path, caption_text)
                entry["insert_status"] = "placeholder"
                entry["insert_location"] = placeholder
                log.write(f"Inserted {entry['id']} at placeholder {placeholder}")
                inserted = True
                break

        if not inserted:
            target_texts = [str(entry.get("insert_after_text") or ""), str(entry.get("insert_before_heading") or "")]
            for target_text in [text for text in target_texts if text]:
                for paragraph in iter_paragraphs(document):
                    if target_text.lower() in paragraph.text.lower():
                        insert_image_in_paragraph(paragraph, image_path, caption_text)
                        entry["insert_status"] = "plan_fallback"
                        entry["insert_location"] = target_text
                        log.write(f"Inserted {entry['id']} near image_plan text: {target_text[:120]}")
                        inserted = True
                        break
                if inserted:
                    break

        if not inserted:
            if not appended_heading_added:
                document.add_heading("Generated Images", level=2)
                appended_heading_added = True
            append_generated_image(document, image_path, caption_text)
            entry["insert_status"] = "appended"
            entry["insert_location"] = "Generated Images"
            log.write(f"Could not find slot for {entry['id']}; appended to Generated Images.")

    remove_leftover_placeholders(document)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)
    log.write(f"Final Word output saved: {output_docx}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Run phase6 image generation and Word insertion.")
    parser.add_argument("--config", type=Path, required=True)
    parser.add_argument("--project-dir", type=Path, required=True)
    parser.add_argument("--final-docx", type=Path, required=True)
    parser.add_argument("--article-md", type=Path, required=True)
    parser.add_argument("--image-plan", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--mode", choices=["all", "generate", "insert"], default="all")
    parser.add_argument("--only-image-id", default="")
    parser.add_argument("--review-comment", default="")
    args = parser.parse_args()

    output_dir = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    log = Phase6Log(output_dir / "phase6.log")
    output_docx = output_dir / "final_article_with_images.docx"
    metadata_path = output_dir / "image_metadata.json"
    review_path = output_dir / "image_review.json"
    compliance_path = output_dir / "image_compliance_report.json"
    images_dir = output_dir / "images"

    try:
        config = read_json(args.config, {})
        log.write("phase6 started")
        log.write(f"mode={args.mode}")
        if args.only_image_id:
            log.write(f"only_image_id={args.only_image_id}")
        log.write(f"image_plan={args.image_plan}")
        log.write(f"slot_article_md={args.article_md}")
        model_fields = resolve_model_request_fields(config)
        log_model_request_fields(log, model_fields)

        if not config.get("enabled", True):
            write_json(metadata_path, {"images": [], "status": "skipped", "reason": "image generation disabled"})
            write_json(compliance_path, {"images": [], "status": "skipped", "reason": "image generation disabled"})
            if args.mode == "all":
                copy_original_docx(args.final_docx, output_docx, log, "image generation disabled")
            else:
                remove_stale_output_docx(output_docx, output_dir, log, "image generation disabled")
            log.write("phase6 completed with image generation disabled")
            return 0

        if args.mode != "insert":
            try:
                validate_model_request_fields(model_fields)
            except ImageModelConfigurationError as exc:
                reason = str(exc)
                write_json(metadata_path, {"images": [], "status": "failed", "error": reason})
                write_json(compliance_path, {"images": [], "status": "failed", "error": reason})
                if args.mode == "all":
                    copy_original_docx(args.final_docx, output_docx, log, reason)
                else:
                    remove_stale_output_docx(output_docx, output_dir, log, reason)
                log.write(reason)
                return 0
            normalize_config_model_fields(config, model_fields)
            write_json(args.config, config)

        if not args.image_plan.exists():
            reason = "image_plan.json missing; phase6 will not generate fallback prompts"
            write_json(metadata_path, {"images": [], "status": "skipped", "reason": reason})
            write_json(compliance_path, {"images": [], "status": "skipped", "reason": reason})
            if args.mode == "all":
                copy_original_docx(args.final_docx, output_docx, log, reason)
            else:
                remove_stale_output_docx(output_docx, output_dir, log, reason)
            log.write(reason)
            return 0

        plan = read_json(args.image_plan, {})
        article_for_count = args.article_md.read_text(encoding="utf-8-sig", errors="replace") if args.article_md.exists() else ""
        plan = ensure_plan_image_count(plan, article_for_count, config, log)
        write_json(args.image_plan, plan)
        ready, reason = policy_ready(plan)
        if not ready:
            write_json(metadata_path, {"images": [], "status": "skipped", "reason": reason})
            write_json(compliance_path, {"images": [], "status": "skipped", "reason": reason})
            if args.mode == "all":
                copy_original_docx(args.final_docx, output_docx, log, reason)
            else:
                remove_stale_output_docx(output_docx, output_dir, log, reason)
            log.write(reason)
            return 0

        valid_images = validate_plan_images(plan, log)
        log.write(f"valid planned images: {len(valid_images)}")
        if not valid_images:
            reason = "image_plan.json contained no valid planned images"
            write_json(metadata_path, {"images": [], "status": "failed", "error": reason})
            write_json(compliance_path, {"images": [], "status": "failed", "error": reason})
            if args.mode == "all":
                copy_original_docx(args.final_docx, output_docx, log, reason)
            else:
                remove_stale_output_docx(output_docx, output_dir, log, reason)
            return 0

        existing_metadata_raw = read_json(metadata_path, {})
        existing_metadata = existing_metadata_raw.get("images", []) if isinstance(existing_metadata_raw, dict) else []
        if not isinstance(existing_metadata, list):
            existing_metadata = []

        if args.only_image_id:
            valid_image_ids = {str(image.get("id") or "") for image in valid_images}
            if args.only_image_id not in valid_image_ids:
                reason = f"Requested image id was not found in image_plan.json: {args.only_image_id}"
                write_json(metadata_path, {"images": existing_metadata, "status": "failed", "error": reason})
                log.write(reason)
                return 0

        if args.mode == "insert":
            metadata = existing_metadata
            compliance_raw = read_json(compliance_path, {})
            compliance_report = compliance_raw.get("images", []) if isinstance(compliance_raw, dict) else []
            if not isinstance(compliance_report, list):
                compliance_report = []
        else:
            selected_images = valid_images
            force_ids: set[str] = set()
            if args.only_image_id:
                selected_images = [image for image in valid_images if str(image.get("id") or "") == args.only_image_id]
                force_ids = {args.only_image_id}
            metadata, compliance_report = generate_images(
                selected_images,
                plan,
                config,
                images_dir,
                log,
                existing_metadata,
                force_regenerate_ids=force_ids,
                review_comment=args.review_comment,
            )
            if args.only_image_id:
                replaced_ids = {str(entry.get("id") or "") for entry in metadata}
                preserved = [
                    entry
                    for entry in existing_metadata
                    if str(entry.get("id") or "") not in replaced_ids
                ]
                metadata = preserved + metadata
                metadata.sort(key=lambda entry: str(entry.get("id") or ""))

        approved_image_ids = approved_image_ids_from_review(review_path)
        planned_count = len(valid_images)
        success_count = (
            insertable_generated_count(metadata, approved_image_ids)
            if args.mode == "insert"
            else successful_generated_count(metadata)
        )
        failed_count = max(0, planned_count - success_count)
        if success_count < planned_count:
            reason = (
                f"Phase 6 generated {success_count}/{planned_count} planned images. "
                "Final Word output will not be generated until every planned image succeeds. "
                "Retry Phase 6 to generate only failed or missing images."
            )
            write_json(
                metadata_path,
                {
                    "images": metadata,
                    "status": "failed",
                    "error": reason,
                    "planned_count": planned_count,
                    "success_count": success_count,
                    "failed_count": failed_count,
                    "retry_needed": True,
                },
            )
            write_json(
                compliance_path,
                {
                    "images": compliance_report,
                    "status": "failed",
                    "error": reason,
                    "planned_count": planned_count,
                    "success_count": success_count,
                    "failed_count": failed_count,
                    "retry_needed": True,
                },
            )
            remove_stale_output_docx(output_docx, output_dir, log, reason)
            log.write(reason)
            return 0

        if args.mode == "generate":
            write_json(
                metadata_path,
                {
                    "images": metadata,
                    "status": "waiting_image_review",
                    "planned_count": planned_count,
                    "success_count": success_count,
                    "failed_count": 0,
                    "retry_needed": False,
                    "review_required": True,
                },
            )
            write_json(
                compliance_path,
                {
                    "images": compliance_report,
                    "status": "waiting_image_review",
                    "planned_count": planned_count,
                    "success_count": success_count,
                    "failed_count": 0,
                    "retry_needed": False,
                    "review_required": True,
                },
            )
            remove_stale_output_docx(output_docx, output_dir, log, "waiting for human image review")
            log.write("phase6 generated images and is waiting for human review")
            return 0

        source_docx = convert_slots_markdown_to_docx(args.article_md, args.final_docx, output_dir, log)
        word_insert_failed = False
        word_insert_error = ""
        try:
            insert_images_into_docx(
                source_docx=source_docx,
                original_docx=args.final_docx,
                output_docx=output_docx,
                metadata=metadata,
                allow_non_compliant=bool(config.get("allow_non_compliant_images", False)),
                approved_image_ids=approved_image_ids,
                log=log,
            )
        except Exception as exc:
            word_insert_failed = True
            word_insert_error = str(exc)
            log.write(f"Word insertion failed: {exc}")
            remove_stale_output_docx(output_docx, output_dir, log, "Word insertion failed")
            for entry in metadata:
                if entry.get("status") == "success":
                    entry["insert_status"] = "failed"
                    entry["insert_error"] = str(exc)

        if word_insert_failed:
            reason = f"Word insertion failed after all planned images were generated: {word_insert_error}"
            write_json(
                metadata_path,
                {
                    "images": metadata,
                    "status": "failed",
                    "error": reason,
                    "planned_count": planned_count,
                    "success_count": success_count,
                    "failed_count": 0,
                    "retry_needed": True,
                },
            )
            write_json(
                compliance_path,
                {
                    "images": compliance_report,
                    "status": "failed",
                    "error": reason,
                    "planned_count": planned_count,
                    "success_count": success_count,
                    "failed_count": 0,
                    "retry_needed": True,
                },
            )
            log.write(reason)
            return 0

        write_json(
            metadata_path,
            {
                "images": metadata,
                "status": "completed",
                "planned_count": planned_count,
                "success_count": success_count,
                "failed_count": 0,
                "retry_needed": False,
            },
        )
        write_json(
            compliance_path,
            {
                "images": compliance_report,
                "status": "completed",
                "planned_count": planned_count,
                "success_count": success_count,
                "failed_count": 0,
                "retry_needed": False,
            },
        )
        log.write("phase6 completed")
        return 0
    except Exception as exc:
        error_message = normalize_image_generation_error(exc)
        log.write(f"phase6 unexpected error: {error_message}")
        write_json(metadata_path, {"images": [], "status": "failed", "error": error_message})
        write_json(compliance_path, {"images": [], "status": "failed", "error": error_message})
        if args.mode == "all":
            copy_original_docx(args.final_docx, output_docx, log, "unexpected phase6 error")
        else:
            remove_stale_output_docx(output_docx, output_dir, log, "unexpected phase6 error")
        return 0
    finally:
        log.write(f"final_output={output_docx}")
        log.save()


if __name__ == "__main__":
    sys.exit(main())
