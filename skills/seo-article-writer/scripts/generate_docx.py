#!/usr/bin/env python3
"""Convert a simple markdown SEO article into a Google Docs-friendly docx."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

IMAGE_SLOT_RE = re.compile(r"^\[IMAGE_\d+\]\s*$")
INTERNAL_LINK_RE = re.compile(r"\[([^\]]+)\]\(INTERNAL_LINK:([^)]+)\)")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a clickable Word hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


HEADING_RE = re.compile(r"^(#{1,3})\s+(.+?)\s*$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.+)")
METADATA_RE = re.compile(r"^(?:\*\*)?(SEO Title|Description|URL):(?:\*\*)?\s+(.+?)\s*$", re.IGNORECASE)
IMAGE_RE = re.compile(r"!\[[^\]]*\]\([^)]*\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
CODE_SPAN_RE = re.compile(r"`([^`]+)`")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def load_image_plan(plan_path: Path | None) -> dict[str, dict[str, str]]:
    """Load image plan and return a dict keyed by [IMAGE_N] marker."""
    if not plan_path or not plan_path.exists():
        return {}
    try:
        data = json.loads(plan_path.read_text(encoding="utf-8-sig"))
    except Exception:
        return {}
    images = data.get("images") if isinstance(data.get("images"), list) else []
    plan: dict[str, dict[str, str]] = {}
    for img in images:
        if not isinstance(img, dict):
            continue
        image_id = str(img.get("id") or "").strip()
        if not image_id:
            continue
        marker = f"[{image_id}]"
        plan[marker] = {
            "id": image_id,
            "type": str(img.get("type") or "").strip(),
            "alt_text": str(img.get("alt_text") or "").strip(),
            "prompt": str(img.get("prompt") or "").strip(),
        }
    return plan


def add_image_placeholder(document: Document, info: dict[str, str]) -> None:
    """Insert a formatted image placeholder paragraph."""
    image_id = info.get("id", "")
    image_type = info.get("type", "")
    alt_text = info.get("alt_text", "")
    prompt = info.get("prompt", "")

    # Title line
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    icon = title.add_run(f"▶ {image_id}")
    icon.bold = True
    if image_type:
        title.add_run(f" — {image_type}")

    # Alt text line
    if alt_text:
        alt_para = document.add_paragraph()
        alt_para.paragraph_format.space_after = Pt(2)
        label = alt_para.add_run("Alt: ")
        label.bold = True
        label.font.size = Pt(10)
        alt_para.add_run(alt_text)

    # Prompt line
    if prompt:
        prompt_para = document.add_paragraph()
        prompt_para.paragraph_format.space_after = Pt(8)
        label = prompt_para.add_run("Prompt: ")
        label.bold = True
        label.font.size = Pt(10)
        run = prompt_para.add_run(prompt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_divider(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value


def strip_inline_markdown(text: str) -> str:
    """Remove inline markdown syntax that is not converted to formatting."""
    text = IMAGE_RE.sub("", text)
    text = LINK_RE.sub(r"\1", text)
    text = CODE_SPAN_RE.sub(r"\1", text)
    return text


def add_runs_with_inline_format(paragraph, text: str, base_url: str = "") -> None:
    """Parse inline markdown and add formatted runs to a paragraph."""
    text = IMAGE_RE.sub("", text)

    # Process internal links as hyperlinks
    def replace_internal_link(match):
        anchor = match.group(1)
        slug = match.group(2)
        if base_url:
            url = base_url.rstrip("/") + "/" + slug.lstrip("/")
            add_hyperlink(paragraph, anchor, url)
        else:
            paragraph.add_run(anchor)
        return ""

    # Process regular markdown links as hyperlinks
    def replace_link(match):
        anchor = match.group(1)
        url = match.group(2)
        if url.startswith("/") and base_url:
            url = base_url.rstrip("/") + url
        if url.startswith("http"):
            add_hyperlink(paragraph, anchor, url)
        else:
            paragraph.add_run(anchor)
        return ""

    # First handle internal links, then regular links
    # We need to process links individually since add_hyperlink modifies the paragraph
    last_end = 0
    for match in INTERNAL_LINK_RE.finditer(text):
        # Add text before the link
        before = text[last_end:match.start()]
        if before:
            _add_formatted_runs(paragraph, before)
        # Add the hyperlink
        anchor = match.group(1)
        slug = match.group(2)
        if base_url:
            url = base_url.rstrip("/") + "/" + slug.lstrip("/")
            add_hyperlink(paragraph, anchor, url)
        else:
            paragraph.add_run(anchor)
        last_end = match.end()

    # Process remaining text (after last internal link)
    remaining = text[last_end:]
    # Now handle regular links in remaining text
    regular_last_end = 0
    for match in LINK_RE.finditer(remaining):
        before = remaining[regular_last_end:match.start()]
        if before:
            _add_formatted_runs(paragraph, before)
        anchor = match.group(1)
        url = match.group(2)
        if url.startswith("/") and base_url:
            url = base_url.rstrip("/") + url
        if url.startswith("http"):
            add_hyperlink(paragraph, anchor, url)
        else:
            paragraph.add_run(anchor)
        regular_last_end = match.end()

    final_remaining = remaining[regular_last_end:]
    if final_remaining:
        _add_formatted_runs(paragraph, final_remaining)


def _add_formatted_runs(paragraph, text: str) -> None:
    """Add bold/code/italic formatted runs (no links)."""
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        bold_match = BOLD_RE.match(part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            continue
        code_match = CODE_SPAN_RE.match(part)
        if code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue
        italic_match = ITALIC_RE.match(part)
        if italic_match:
            run = paragraph.add_run(italic_match.group(1))
            run.italic = True
            continue
        paragraph.add_run(part)


def add_paragraph(document: Document, line: str, base_url: str = "") -> None:
    metadata = METADATA_RE.match(line)
    if metadata:
        paragraph = document.add_paragraph()
        label = paragraph.add_run(f"{metadata.group(1)}: ")
        label.bold = True
        add_runs_with_inline_format(paragraph, metadata.group(2).strip(), base_url)
        paragraph.paragraph_format.space_after = Pt(4)
        return

    heading = HEADING_RE.match(line)
    if heading:
        level = len(heading.group(1))
        heading_text = strip_inline_markdown(heading.group(2))
        document.add_heading(heading_text, level=level)
        return

    if line.startswith("- ") or line.startswith("* "):
        bullet_text = line[2:].strip()
        if BOLD_RE.search(bullet_text) or CODE_SPAN_RE.search(bullet_text) or IMAGE_RE.search(bullet_text) or LINK_RE.search(bullet_text):
            paragraph = document.add_paragraph(style="List Bullet")
            add_runs_with_inline_format(paragraph, bullet_text, base_url)
            paragraph.paragraph_format.space_after = Pt(4)
        else:
            document.add_paragraph(bullet_text, style="List Bullet")
        return

    ordered = ORDERED_RE.match(line)
    if ordered:
        num_text = ordered.group(1).strip()
        if BOLD_RE.search(num_text) or CODE_SPAN_RE.search(num_text) or LINK_RE.search(num_text):
            paragraph = document.add_paragraph(style="List Number")
            add_runs_with_inline_format(paragraph, num_text, base_url)
            paragraph.paragraph_format.space_after = Pt(4)
        else:
            document.add_paragraph(num_text, style="List Number")
        return

    if BOLD_RE.search(line) or CODE_SPAN_RE.search(line) or IMAGE_RE.search(line) or LINK_RE.search(line):
        paragraph = document.add_paragraph()
        add_runs_with_inline_format(paragraph, line, base_url)
        paragraph.paragraph_format.space_after = Pt(8)
    else:
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.space_after = Pt(8)


def convert_markdown(markdown: str, image_plan: dict[str, dict[str, str]] | None = None, base_url: str = "") -> Document:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    table_rows: list[list[str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            if table_rows:
                add_table(document, table_rows)
                table_rows = []
            continue

        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            if not is_table_divider(line):
                table_rows.append(split_table_row(line))
            continue

        if table_rows:
            add_table(document, table_rows)
            table_rows = []

        if image_plan and IMAGE_SLOT_RE.match(line.strip()):
            info = image_plan.get(line.strip())
            if info:
                add_image_placeholder(document, info)
                continue

        add_paragraph(document, line, base_url)

    if table_rows:
        add_table(document, table_rows)

    return document


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert markdown article to docx.")
    parser.add_argument("input_md", type=Path, help="Input markdown article path.")
    parser.add_argument("output_docx", type=Path, help="Output docx path.")
    parser.add_argument("--image-plan", type=Path, default=None, help="Optional image_plan.json for placeholder insertion.")
    parser.add_argument("--base-url", default="", help="Base URL for resolving relative links (e.g. https://www.hellotalk.com/en/blog)")
    args = parser.parse_args()

    markdown = args.input_md.read_text(encoding="utf-8")
    image_plan = load_image_plan(args.image_plan)
    document = convert_markdown(markdown, image_plan, args.base_url)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_docx)
    print(f"Wrote {args.output_docx}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
